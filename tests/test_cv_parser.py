"""Tests for CV parser module."""

import json
from pathlib import Path

import pytest

from src.cv_parser import (
    extract_text,
    extract_profile,
    save_profile,
    load_profile,
    _match_terms,
    _find_skills_in_text,
    _find_job_titles,
    _find_locations,
    _categorise_skills,
    _discover_freeform_skills,
    _extract_skills_section_text,
    _is_likely_skill,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def sample_docx(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("AI Engineer with Python, PyTorch, TensorFlow experience.")
    doc.add_paragraph("Based in London, UK. Worked with AWS, Docker, LangChain.")
    path = tmp_path / "test_cv.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def sample_pdf(tmp_path):
    """Create a minimal valid PDF with extractable text."""
    # Minimal hand-crafted PDF — avoids fpdf2/cryptography dependency
    text_content = "AI Engineer with Python, PyTorch, TensorFlow experience. Based in London, UK. Worked with AWS, Docker, LangChain."
    stream = (
        f"BT /F1 12 Tf 100 700 Td ({text_content}) Tj ET"
    )
    stream_bytes = stream.encode("latin-1")
    objects = []
    # obj 1: catalog
    objects.append(b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n")
    # obj 2: pages
    objects.append(b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n")
    # obj 3: page
    objects.append(
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
    )
    # obj 4: content stream
    objects.append(
        f"4 0 obj\n<< /Length {len(stream_bytes)} >>\nstream\n".encode("latin-1")
        + stream_bytes
        + b"\nendstream\nendobj\n"
    )
    # obj 5: font
    objects.append(
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
    )

    body = b"%PDF-1.4\n"
    offsets = []
    for obj in objects:
        offsets.append(len(body))
        body += obj

    xref_offset = len(body)
    xref = f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n"
    for off in offsets:
        xref += f"{off:010d} 00000 n \n"
    xref += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    )
    body += xref.encode("latin-1")

    path = tmp_path / "test_cv.pdf"
    path.write_bytes(body)
    return path


# ---------------------------------------------------------------------------
# Text extraction tests
# ---------------------------------------------------------------------------

def test_extract_text_docx(sample_docx):
    text = extract_text(str(sample_docx))
    assert "Python" in text
    assert "PyTorch" in text
    assert "London" in text


def test_extract_text_pdf(sample_pdf):
    text = extract_text(str(sample_pdf))
    assert "Python" in text
    assert "PyTorch" in text
    assert "London" in text


def test_extract_text_unsupported(tmp_path):
    txt_file = tmp_path / "resume.txt"
    txt_file.write_text("some text")
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(str(txt_file))


def test_extract_text_missing_file():
    with pytest.raises(FileNotFoundError):
        extract_text("/nonexistent/path/cv.pdf")


# ---------------------------------------------------------------------------
# Legacy term matching tests (backward compat)
# ---------------------------------------------------------------------------

def test_match_terms_finds_skills():
    matches = _match_terms("Python and PyTorch developer", ["Python", "Java"])
    assert matches == ["Python"]


def test_match_terms_case_insensitive():
    matches = _match_terms("python pytorch tensorflow", ["Python", "PyTorch"])
    assert "Python" in matches
    assert "PyTorch" in matches


def test_match_terms_no_match():
    matches = _match_terms("marketing SEO analytics", ["Python", "PyTorch"])
    assert matches == []


def test_match_terms_substring_matching():
    matches = _match_terms("Experience with Deep Learning models", ["Deep Learning"])
    assert "Deep Learning" in matches


# ---------------------------------------------------------------------------
# Smart extraction tests
# ---------------------------------------------------------------------------

def test_find_skills_discovers_python():
    counts = _find_skills_in_text("Python developer using Python and Django")
    assert "Python" in counts
    assert counts["Python"] >= 2


def test_find_skills_discovers_java():
    """A Java developer's CV should find Java, not just AI/ML skills."""
    counts = _find_skills_in_text(
        "Java Spring Boot developer with MySQL and Docker experience. "
        "Worked on microservices using Java and Kubernetes."
    )
    assert "Java" in counts
    assert "Spring Boot" in counts
    assert "MySQL" in counts
    assert "Docker" in counts


def test_find_skills_discovers_frontend():
    counts = _find_skills_in_text(
        "React TypeScript developer. Built apps with Next.js, Tailwind CSS, Redux."
    )
    assert "React" in counts
    assert "TypeScript" in counts
    assert "Next.js" in counts


def test_find_job_titles_multi_domain():
    text = "Worked as a Software Engineer at Google, then Data Scientist at Meta."
    titles = _find_job_titles(text)
    assert "Software Engineer" in titles
    assert "Data Scientist" in titles


def test_find_locations_multiple():
    text = "Based in London, open to Remote and Hybrid roles."
    locs = _find_locations(text)
    assert "London" in locs
    assert "Remote" in locs


def test_find_locations_uk():
    text = "Worked in Manchester, England. Relocated to Edinburgh, Scotland."
    locs = _find_locations(text)
    assert "Manchester" in locs
    assert "England" in locs
    assert "Edinburgh" in locs


# ---------------------------------------------------------------------------
# Freeform discovery tests — skills NOT in the database
# ---------------------------------------------------------------------------

def test_discover_freeform_from_skills_section():
    """Skills section items not in KNOWN_SKILLS should be discovered."""
    text = (
        "Experience\nBuilt systems for clients.\n\n"
        "Skills:\n"
        "Pulumi, Temporal, dbt Cloud, Prefect, Airbyte\n\n"
        "Education\nBSc Computer Science"
    )
    discovered = _discover_freeform_skills(text)
    # Pulumi and Temporal are NOT in KNOWN_SKILLS — should be discovered
    # (dbt and Prefect ARE in KNOWN_SKILLS, so they won't appear here)
    assert "Pulumi" in discovered or "Temporal" in discovered


def test_discover_freeform_from_context_pattern():
    """'experience with X, Y' pattern should capture unknown skills."""
    text = "I have experience with Mage.ai, ZenML, and Metaflow for ML pipelines."
    discovered = _discover_freeform_skills(text)
    # These are niche tools unlikely to be in KNOWN_SKILLS
    found_any = any(
        term in discovered
        for term in ["Mage.ai", "ZenML", "Metaflow"]
    )
    assert found_any


def test_discover_freeform_skips_noise():
    """Noise words and short terms should not be discovered."""
    text = (
        "Skills:\n"
        "the, and, or, a, in, for, with, Python\n"
    )
    discovered = _discover_freeform_skills(text)
    assert "the" not in discovered
    assert "and" not in discovered
    # Python is in KNOWN_SKILLS, so it won't be in freeform either
    assert "Python" not in discovered


def test_is_likely_skill_filters():
    assert _is_likely_skill("Kubernetes") is True
    assert _is_likely_skill("ArgoCD") is True
    assert _is_likely_skill("the") is False
    assert _is_likely_skill("and") is False
    assert _is_likely_skill("") is False
    assert _is_likely_skill("a") is False
    assert _is_likely_skill("12345") is False
    # Too many words
    assert _is_likely_skill("this is a very long sentence about things") is False


def test_extract_skills_section_text():
    text = (
        "Summary\nI am a developer.\n\n"
        "Technical Skills:\n"
        "Python, Rust, ArgoCD, Pulumi\n\n"
        "Experience\nWorked at Google."
    )
    sections = _extract_skills_section_text(text)
    assert len(sections) >= 1
    combined = " ".join(sections).lower()
    assert "pulumi" in combined


def test_freeform_skills_end_up_in_profile():
    """Freeform-discovered skills should appear in the final profile."""
    text = (
        "Software Engineer\n\n"
        "Skills:\n"
        "Python, Java, Pulumi, Temporal, ZenML, ClearML\n\n"
        "Experience:\n"
        "Built infrastructure with Pulumi and Temporal. "
        "Used ZenML and ClearML for ML experiment tracking.\n"
        "Based in London."
    )
    profile = extract_profile(text)
    all_skills = (
        profile["primary_skills"]
        + profile["secondary_skills"]
        + profile["tertiary_skills"]
    )
    # Python and Java are in KNOWN_SKILLS
    assert "Python" in all_skills
    assert "Java" in all_skills
    # Pulumi and Temporal are NOT in KNOWN_SKILLS but should be discovered
    # (they appear in Skills section + experience = should be primary or secondary)
    freeform_found = [s for s in all_skills if s in ("Pulumi", "Temporal", "ZenML", "ClearML")]
    assert len(freeform_found) >= 1, f"Expected freeform skills, got: {all_skills}"


def test_freeform_skills_used_in_relevance_keywords():
    """Freeform-discovered skills should appear in relevance keywords."""
    import src.cv_parser as cv_mod
    import src.filters.skill_matcher as sm
    import json
    from pathlib import Path
    import tempfile

    text = (
        "Skills:\n"
        "Python, Pulumi, Temporal, ZenML\n\n"
        "Experience with Pulumi for IaC. Used Temporal for workflows."
    )
    profile = extract_profile(text)
    all_skills = (
        profile["primary_skills"]
        + profile["secondary_skills"]
        + profile["tertiary_skills"]
    )
    # Check that at least one freeform skill is captured
    freeform_in_profile = [s for s in all_skills if s in ("Pulumi", "Temporal", "ZenML")]
    assert len(freeform_in_profile) >= 1


# ---------------------------------------------------------------------------
# Auto-categorisation tests
# ---------------------------------------------------------------------------

def test_categorise_skills_by_frequency():
    """Skills mentioned more often should rank higher."""
    from collections import Counter
    counts = Counter({"Python": 5, "Docker": 2, "Nginx": 1})
    text = "Python Python Python Python Python Docker Docker Nginx"
    primary, secondary, tertiary = _categorise_skills(counts, text)
    assert "Python" in primary
    assert "Nginx" in tertiary


def test_categorise_skills_section_boost():
    """Skills in a 'Skills' section get boosted."""
    from collections import Counter
    text = "Experience with various tools.\n\nSkills:\nReact, Node.js, PostgreSQL\n\nWork history..."
    counts = _find_skills_in_text(text)
    primary, secondary, tertiary = _categorise_skills(counts, text)
    # React appears in skills section (once) → should be at least secondary
    all_skills = primary + secondary
    assert "React" in all_skills or "Node.js" in all_skills


# ---------------------------------------------------------------------------
# Profile extraction tests (integration)
# ---------------------------------------------------------------------------

def test_extract_profile_ai_engineer():
    """AI/ML CV should produce AI-relevant profile."""
    text = (
        "AI Engineer with Python, PyTorch, TensorFlow experience. "
        "Expert in LangChain, RAG, LLM, NLP, Deep Learning. "
        "Used AWS, Docker, Kubernetes, FastAPI. "
        "Based in London, UK."
    )
    profile = extract_profile(text)
    all_skills = (
        profile["primary_skills"]
        + profile["secondary_skills"]
        + profile["tertiary_skills"]
    )
    assert "Python" in all_skills
    assert "PyTorch" in all_skills
    assert "AWS" in all_skills
    assert "London" in profile["locations"]
    assert "AI Engineer" in profile["job_titles"]


def test_extract_profile_java_developer():
    """A Java developer's CV should produce Java-relevant profile — NOT empty."""
    text = (
        "Software Engineer\n\n"
        "Skills:\n"
        "Java, Spring Boot, MySQL, PostgreSQL, Docker, Kubernetes, "
        "React, TypeScript, Git, Jenkins, AWS\n\n"
        "Experience:\n"
        "Built microservices in Java using Spring Boot. Deployed with Docker "
        "and Kubernetes on AWS. Frontend in React and TypeScript.\n\n"
        "Location: Manchester, UK"
    )
    profile = extract_profile(text)
    all_skills = (
        profile["primary_skills"]
        + profile["secondary_skills"]
        + profile["tertiary_skills"]
    )
    assert "Java" in all_skills
    assert "Spring Boot" in all_skills
    assert "Docker" in all_skills
    assert "Software Engineer" in profile["job_titles"]
    assert "Manchester" in profile["locations"]
    # Python should NOT appear — it's not in this CV
    assert "Python" not in all_skills


def test_extract_profile_accountant():
    """Non-tech CV should still extract meaningful info."""
    text = (
        "Accountant with 10 years experience.\n"
        "Proficient in Excel, QuickBooks, Xero, Sage.\n"
        "Based in Birmingham, UK."
    )
    profile = extract_profile(text)
    all_skills = (
        profile["primary_skills"]
        + profile["secondary_skills"]
        + profile["tertiary_skills"]
    )
    assert "Excel" in all_skills
    assert "Birmingham" in profile["locations"]
    assert "Accountant" in profile["job_titles"]


def test_extract_profile_empty_text():
    profile = extract_profile("")
    assert profile["job_titles"] == []
    assert profile["primary_skills"] == []
    assert profile["secondary_skills"] == []
    assert profile["tertiary_skills"] == []
    assert profile["locations"] == []


def test_extract_profile_has_metadata():
    profile = extract_profile("Python developer")
    assert "extracted_at" in profile
    # Should be a valid ISO timestamp
    from datetime import datetime
    datetime.fromisoformat(profile["extracted_at"])


def test_different_cvs_produce_different_profiles():
    """The core multi-user requirement: different CVs → different results."""
    ai_cv = (
        "AI Engineer with Python, PyTorch, TensorFlow, LangChain, RAG. "
        "Based in London."
    )
    java_cv = (
        "Software Engineer with Java, Spring Boot, MySQL, React. "
        "Based in Manchester."
    )

    ai_profile = extract_profile(ai_cv)
    java_profile = extract_profile(java_cv)

    ai_skills = set(
        ai_profile["primary_skills"]
        + ai_profile["secondary_skills"]
        + ai_profile["tertiary_skills"]
    )
    java_skills = set(
        java_profile["primary_skills"]
        + java_profile["secondary_skills"]
        + java_profile["tertiary_skills"]
    )

    # They should have different skill sets
    assert ai_skills != java_skills
    assert "PyTorch" in ai_skills
    assert "PyTorch" not in java_skills
    assert "Spring Boot" in java_skills
    assert "Spring Boot" not in ai_skills


# ---------------------------------------------------------------------------
# Save / Load tests
# ---------------------------------------------------------------------------

def test_save_and_load_roundtrip(tmp_path):
    profile = {
        "job_titles": ["AI Engineer"],
        "primary_skills": ["Python", "PyTorch"],
        "secondary_skills": ["AWS"],
        "tertiary_skills": ["Git"],
        "locations": ["London"],
        "source_file": "cv.pdf",
        "extracted_at": "2025-01-01T00:00:00+00:00",
    }
    path = tmp_path / "profile.json"
    save_profile(profile, path)
    loaded = load_profile(path)
    assert loaded == profile


def test_load_missing_file(tmp_path):
    result = load_profile(tmp_path / "nonexistent.json")
    assert result is None


def test_save_creates_parent_dirs(tmp_path):
    profile = {"job_titles": [], "primary_skills": []}
    path = tmp_path / "nested" / "deep" / "profile.json"
    save_profile(profile, path)
    assert path.exists()


# ---------------------------------------------------------------------------
# Integration with skill_matcher
# ---------------------------------------------------------------------------

def test_scorer_uses_cv_profile(tmp_path, monkeypatch):
    """When a CV profile exists, scorer should use its skill lists."""
    from datetime import datetime, timezone
    from src.models import Job
    import src.filters.skill_matcher as sm
    import src.cv_parser as cv_mod

    profile_path = tmp_path / "cv_profile.json"
    # Profile with ONLY "Python" as primary skill
    profile = {
        "job_titles": ["AI Engineer"],
        "primary_skills": ["Python"],
        "secondary_skills": [],
        "tertiary_skills": [],
        "locations": ["London"],
    }
    profile_path.write_text(json.dumps(profile))

    monkeypatch.setattr(cv_mod, "CV_PROFILE_PATH", profile_path)
    sm.reload_profile()

    job = Job(
        title="AI Engineer",
        company="Test",
        apply_url="https://example.com",
        source="test",
        date_found=datetime.now(timezone.utc).isoformat(),
        location="London",
        description="Python developer with AWS and Docker experience",
    )
    score = sm.score_job(job)
    # With only Python as a primary skill, AWS and Docker won't contribute
    # Score should be lower than with the full keyword list
    assert score > 0
    sm.reload_profile()


def test_scorer_falls_back_without_profile(tmp_path, monkeypatch):
    """Without a CV profile, scorer should use default keywords.py lists."""
    from datetime import datetime, timezone
    from src.models import Job
    import src.filters.skill_matcher as sm
    import src.cv_parser as cv_mod

    monkeypatch.setattr(cv_mod, "CV_PROFILE_PATH", tmp_path / "nonexistent.json")
    sm.reload_profile()

    job = Job(
        title="AI Engineer",
        company="Test",
        apply_url="https://example.com",
        source="test",
        date_found=datetime.now(timezone.utc).isoformat(),
        location="London, UK",
        description=(
            "Python PyTorch TensorFlow LangChain RAG LLM NLP "
            "Deep Learning AWS Docker Kubernetes"
        ),
    )
    score = sm.score_job(job)
    # Should score high with the full default list
    assert score >= 70
    sm.reload_profile()


def test_scorer_explicit_profile_parameter():
    """score_job should accept an explicit profile, bypassing the global cache."""
    from datetime import datetime, timezone
    from src.models import Job
    from src.filters.skill_matcher import score_job

    java_profile = {
        "job_titles": ["Software Engineer"],
        "primary_skills": ["Java", "Spring Boot"],
        "secondary_skills": ["MySQL", "Docker"],
        "tertiary_skills": ["Git"],
        "locations": ["Manchester"],
    }

    java_job = Job(
        title="Software Engineer",
        company="Test",
        apply_url="https://example.com",
        source="test",
        date_found=datetime.now(timezone.utc).isoformat(),
        location="Manchester",
        description="Java Spring Boot developer with MySQL and Docker",
    )

    ai_job = Job(
        title="AI Engineer",
        company="Test",
        apply_url="https://example.com",
        source="test",
        date_found=datetime.now(timezone.utc).isoformat(),
        location="London",
        description="Python PyTorch TensorFlow LangChain RAG",
    )

    # Java job should score HIGH with Java profile
    java_score = score_job(java_job, profile=java_profile)
    # AI job should score LOW with Java profile
    ai_score = score_job(ai_job, profile=java_profile)

    assert java_score > ai_score
