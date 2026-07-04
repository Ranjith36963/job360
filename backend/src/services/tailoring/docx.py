"""ATS-friendly DOCX rendering (spec guardrail #5) — sibling of ``pdf.py``.

Same discipline as the PDF path: one column, one standard font, real editable text,
no tables or graphics — so applicant-tracking software parses it cleanly. ``python-docx``
is lazy-imported (heavy-dep discipline, root rules #11/#16).
"""

from __future__ import annotations


def render_docx(text: str, title: str = "") -> bytes:
    """Render plain text to a clean, single-column, machine-readable .docx (bytes)."""
    from io import BytesIO

    from docx import Document  # lazy import
    from docx.shared import Pt

    doc = Document()

    # ATS-friendly base style: single standard font, sensible size.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(14)

    for line in (text or "").split("\n"):
        # Every line becomes its own paragraph so spacing + structure survive; blank
        # lines still add a paragraph so gaps between sections are preserved.
        doc.add_paragraph(line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
